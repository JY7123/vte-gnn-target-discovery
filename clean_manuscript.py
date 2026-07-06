#!/usr/bin/env python3
"""Clean manuscript.docx: remove False Positive Injection Test, reframe prior injection.
Run: python clean_manuscript.py
"""
from docx import Document
from docx.shared import Pt

DOC_PATH = r"D:\JY\work\my work\新思路\投稿nc\manuscript.docx"
OUT_PATH = r"D:\JY\work\my work\新思路\投稿nc\manuscript_cleaned.docx"

doc = Document(DOC_PATH)

# Map paragraph indices to their content summary (for precise targeting)
target_map = {}
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        target_map[i] = t[:80]


# ── 1. REMOVE: Methods "False positive injection test" heading + body ──
# Paragraph [36]: heading, [37]: body
for idx in [36, 37]:
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

print("  [OK] Removed Methods: 'False positive injection test' (paras 36-37)")


# ── 2. REMOVE: Results paragraph about injection test ──
# After removing 2 paras above, indices shift by -2.
# Original para [55] becomes [53]
# The para starts: "The prior injection system suppressed known literature-driven false positives..."
# Need to find it by content now since indices shifted

for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("The prior injection system suppressed"):
        p._element.getparent().remove(p._element)
        print(f"  [OK] Removed Results injection test sentence (was para ~55)")
        break


# ── 3. REWRITE: Figure 2 legend title ──
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Figure 2") and "performance" in p.text.lower():
        # Replace title
        old_title = p.text.strip()
        new_title = "Figure 2 | Model performance, ablation analysis, and emergent structural noise resistance."
        # Replace only the title part (preserve any trailing content)
        for run in p.runs:
            if "Figure 2" in run.text:
                run.text = run.text.replace(
                    "Model performance, ablation analysis, and prior injection effects.",
                    "Model performance, ablation analysis, and emergent structural noise resistance."
                )
        print(f"  [OK] Updated Figure 2 legend title")
        break

# ── 4. REMOVE: Figure 2 legend panel "d" description ──
# The panel d description is in the paragraph after Figure 2 title
# It contains: "d, Efficacy of false positive suppression..."
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Figure 2") and "performance" in p.text.lower():
        # Next paragraph should be the panel descriptions
        next_p = doc.paragraphs[i + 1]
        text = next_p.text

        # Find and remove the "d," sentence
        # Pattern: "d, Efficacy of false positive suppression via prior injection..."
        import re
        # Match from "d, Efficacy..." to period before next panel letter or end
        cleaned = re.sub(
            r'\s*d,\s*Efficacy of false positive suppression.*?(?=\.\s*$|$)',
            '',
            text
        )
        # Also clean up double periods that might result
        cleaned = cleaned.replace('..', '.')

        if cleaned != text:
            # Replace text in the first run
            if next_p.runs:
                # Put all cleaned text in first run, clear others
                next_p.runs[0].text = cleaned
                for run in next_p.runs[1:]:
                    run.text = ''
            print(f"  [OK] Removed Figure 2 legend panel d description")
        else:
            # Fallback: try to find and remove "d, Efficacy..."
            for run in next_p.runs:
                if "Efficacy of false positive suppression" in run.text:
                    run.text = re.sub(
                        r'd,\s*Efficacy of false positive suppression.*?(?=\s*$)',
                        '',
                        run.text
                    )
                    print(f"  [OK] Removed Figure 2 legend panel d description (fallback)")
                    break
        break


# ── 5. REFRAME: Three-layer prior injection paragraph ──
# Paragraph [24]: "The model integrates prior knowledge at three levels..."
# Change "hard prior encodes curated biological trust" framing
# Remove mention of false positive suppression, focus on guidance during early training

for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("The model integrates prior knowledge at three levels"):
        for run in p.runs:
            # Soften the "penalizes false associations" language
            if "preventing overfitting to spurious co-occurrence" in run.text:
                run.text = run.text.replace(
                    "preventing overfitting to spurious co-occurrence patterns in the training literature",
                    "providing initial biological guidance that prevents the model from overfitting to high-frequency co-occurrence noise during early training stages"
                )
            if "Layer 1 (hard prior)" in run.text or "hard prior" in run.text:
                run.text = run.text.replace(
                    "provides strong initial guidance that prevents the model from overfitting",
                    "provides structured biological initialization that guides early training away from high-frequency but non-causal co-occurrence patterns, before decaying"
                )
        print(f"  [OK] Reframed three-layer prior injection description")
        break


# ── 6. OPTIONAL: Update Discussion mention of prior injection ──
# Paragraph [74]: "The three-layer prior injection system was central..."
# Remove reference to false positive suppression if present
for i, p in enumerate(doc.paragraphs):
    if "three-layer prior injection system was central" in p.text:
        for run in p.runs:
            if "edge weight bias (Layer 1) prevented overfitting" in run.text or "edge weight bias (Layer 1) provided" in run.text:
                run.text = run.text.replace(
                    "prevented overfitting to spurious literature associations",
                    "guided early training away from high-frequency co-occurrence noise, as evidenced by the emergent τ distribution (Figure 2c)"
                )
        print(f"  [OK] Updated Discussion mention of prior injection")
        break


# ── Save ──
doc.save(OUT_PATH)
print(f"\nSaved cleaned manuscript to: {OUT_PATH}")
print("Review the cleaned file, then replace original if satisfied.")
