#!/usr/bin/env python3
"""
Build ATVB-formatted Word manuscript from markdown source.

Converts manuscript/manuscript_draft.md → 投稿ATVB/manuscript.docx
with proper gene italicization (human ALL_CAPS, mouse Title_case),
LaTeX math → Unicode, and ATVB-compliant section structure.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "manuscript" / "manuscript_draft.md"
SUPP_PATH = PROJECT_ROOT / "manuscript" / "supplementary_tables.md"
OUT_DIR = PROJECT_ROOT.parent / "投稿ATVB"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "manuscript.docx"

# ── LaTeX → Unicode ────────────────────────────────────────────────
LATEX_TO_UNICODE = {
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\epsilon": "ε", r"\zeta": "ζ", r"\eta": "η", r"\theta": "θ",
    r"\iota": "ι", r"\kappa": "κ", r"\lambda": "λ", r"\mu": "μ",
    r"\nu": "ν", r"\xi": "ξ", r"\pi": "π", r"\rho": "ρ",
    r"\sigma": "σ", r"\tau": "τ", r"\upsilon": "υ", r"\phi": "φ",
    r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
    r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ", r"\Lambda": "Λ",
    r"\Xi": "Ξ", r"\Pi": "Π", r"\Sigma": "Σ", r"\Phi": "Φ",
    r"\Psi": "Ψ", r"\Omega": "Ω",
    r"\le": "≤", r"\ge": "≥", r"\pm": "±", r"\times": "×",
    r"\rightarrow": "→", r"\leftarrow": "←", r"\Rightarrow": "⇒",
    r"\approx": "≈", r"\propto": "∝", r"\infty": "∞",
    r"\sqrt": "√", r"\partial": "∂", r"\cdot": "·",
    r"\ldots": "…", r"\cdots": "⋯",
    r"\text{": "", r"\text ": "",
}

# Mouse gene pattern: Title case with optional numbers, min 3 chars (e.g. Tgfb1, Smad4)
MOUSE_GENE_RE = re.compile(r"^[A-Z][a-z][a-z0-9]{2,}$")

# Human gene pattern: ALL_CAPS at least 2 chars with optional numbers (not inside *...*)
HUMAN_GENE_RE = re.compile(r"\b([A-Z][A-Z0-9]{2,})\b")


def latex_math_to_unicode(text: str) -> str:
    """Convert LaTeX math expressions to Unicode."""
    # Handle subscripts: X_{text} → X_text (simplified for Word)
    text = re.sub(r"_\{\\text\{([^}]+)\}\}", r"_{\1}", text)
    text = re.sub(r"_\{(\\?[A-Za-z0-9]+)\}", r"_{\1}", text)
    # Handle superscripts
    text = re.sub(r"\^\{([^}]+)\}", r"^{\1}", text)
    # Remove remaining \text{...} wrappers
    text = re.sub(r"\\text\{([^}]+)\}", r"\1", text)
    for latex, uni in LATEX_TO_UNICODE.items():
        text = text.replace(latex, uni)
    return text


def is_mouse_gene(text: str) -> bool:
    """Check if text looks like a mouse gene symbol (Title_case)."""
    # Exclude common short words that match the pattern
    COMMON_WORDS = {"The", "In", "To", "We", "Our", "Day", "For", "And",
                    "Not", "But", "All", "Its", "Via", "Has", "Had", "Was",
                    "Are", "Can", "May", "New", "One", "Two", "His", "Her"}
    if text in COMMON_WORDS:
        return False
    return bool(MOUSE_GENE_RE.match(text))


def process_inline_markdown(paragraph, text: str):
    """
    Process inline formatting and add runs to paragraph.
    Handles: **bold**, *italic*, $math/gene$, plain text.
    """
    # First, protect $...$ blocks by replacing them with placeholders
    math_blocks = []

    def save_math(m):
        math_blocks.append(m.group(1))
        return f"\x00MATH{len(math_blocks) - 1}\x00"
    text = re.sub(r"\$([^$]+?)\$", save_math, text)

    # Split remaining text into formatting segments
    # Handle **bold** and *italic*
    segments = []
    pos = 0
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")

    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            segments.append(("plain", text[pos:m.start()]))
        if m.group(2):  # **bold**
            segments.append(("bold", m.group(2)))
        elif m.group(3):  # *italic*
            segments.append(("italic", m.group(3)))
        pos = m.end()

    # Remaining text
    if pos < len(text):
        segments.append(("plain", text[pos:]))

    # If no segments, treat whole text as plain
    if not segments:
        segments = [("plain", text)]

    # Build runs
    for fmt, seg_text in segments:
        # Restore math blocks and process
        parts = re.split(r"\x00MATH(\d+)\x00", seg_text)
        for i, part in enumerate(parts):
            if not part:
                continue
            if i % 2 == 0:
                # Plain text part
                if part.strip():
                    run = paragraph.add_run(part)
            else:
                # Math block
                math_idx = int(part)
                math_content = math_blocks[math_idx]

                # Always convert LaTeX to Unicode in math blocks
                converted = latex_math_to_unicode(math_content)

                if is_mouse_gene(converted):
                    # Mouse gene → italic (e.g. Tgfb1, Smad4)
                    run = paragraph.add_run(converted)
                    run.italic = True
                elif HUMAN_GENE_RE.match(converted) and not converted[0].isdigit():
                    # Human gene in math mode → italic
                    run = paragraph.add_run(converted)
                    run.italic = True
                else:
                    run = paragraph.add_run(converted)

            if fmt == "bold":
                run.bold = True
            elif fmt == "italic":
                run.italic = True


def add_heading_styled(doc, text: str, level: int):
    """Add a heading with ATVB-appropriate style."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_body_paragraph(doc, text: str):
    """Add a body paragraph with inline formatting."""
    if not text.strip():
        return doc.add_paragraph("")
    p = doc.add_paragraph()
    process_inline_markdown(p, text)
    return p


def add_table_cell_runs(p, text: str, size: int, bold: bool = False):
    """Add runs to paragraph p, honoring **bold** and *italic* inline markers."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            run.bold = bold
            run.font.size = Pt(size)
        if m.group(2):
            run = p.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(size)
        else:
            run = p.add_run(m.group(3))
            run.italic = True
            run.font.size = Pt(size)
        pos = m.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        run.bold = bold
        run.font.size = Pt(size)


def add_table_from_markdown(doc, header_row: list[str], data_rows: list[list[str]]):
    """Create a formatted Word table."""
    ncols = len(header_row)
    nrows = 1 + len(data_rows)
    table = doc.add_table(rows=nrows, cols=ncols, style="Table Grid")

    # Header row
    for j, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_table_cell_runs(p, cell_text, size=9, bold=True)
        # Light gray background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9D9D9")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row in enumerate(data_rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_table_cell_runs(p, cell_text, size=9)

    doc.add_paragraph("")  # spacing after table
    return table


def parse_markdown_table(lines: list[str], start_idx: int) -> tuple:
    """Parse a markdown table starting at start_idx. Returns (header, rows, end_idx)."""
    # Header line
    header = [c.strip() for c in lines[start_idx].strip("|").split("|")]
    # Separator line (skip: start_idx + 1)
    # Data rows
    rows = []
    i = start_idx + 2
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append([c.strip() for c in lines[i].strip("|").split("|")])
        i += 1
    return header, rows, i


def build_manuscript():
    """Main build function."""
    if not MD_PATH.exists():
        print(f"ERROR: {MD_PATH} not found")
        sys.exit(1)

    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # ── Page setup ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Default font ────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    # ── Title Page ──────────────────────────────────────────────
    # ATVB requires: title, authors, affiliations, corresponding author, word count
    i = 0
    title_line = ""
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("# ") and not line.startswith("## "):
            title_line = line[2:]
            break
        i += 1

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title_line)
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Author placeholder
    doc.add_paragraph("")
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run("[AUTHOR LIST — TO BE COMPLETED]").bold = True
    author_p.add_run(f"\n[AFFILIATIONS]")
    author_p.add_run(f"\n\nCorresponding author: [NAME, EMAIL, ADDRESS]")
    author_p.add_run(f"\n\nWord count: [TO BE CALCULATED]")

    doc.add_page_break()

    # ── Structured Abstract ─────────────────────────────────────
    add_heading_styled(doc, "Abstract", level=1)

    # Find abstract content between "## Abstract" and next "## "
    abstract_start = None
    abstract_end = None
    for idx, line in enumerate(lines):
        if line.strip() == "## Abstract" and abstract_start is None:
            abstract_start = idx + 1
        elif abstract_start is not None and line.strip().startswith("## ") and "Abstract" not in line:
            abstract_end = idx
            break
    if abstract_end is None:
        abstract_end = len(lines)

    # Collect abstract paragraphs
    abstract_lines = []
    for idx in range(abstract_start, abstract_end):
        l = lines[idx].strip()
        if l and not l.startswith("---"):
            abstract_lines.append(l)

    # Parse abstract into structured sections
    abstract_text = " ".join(abstract_lines)
    # Split by **...:** markers
    sections = re.split(r"(\*\*[A-Za-z ]+:\*\*)", abstract_text)
    current_heading = "Background:"
    for part in sections:
        part = part.strip()
        if not part:
            continue
        if part.startswith("**") and part.endswith(":**"):
            current_heading = part
            continue
        # Add subsection
        heading_text = current_heading.strip("*:")
        p = doc.add_paragraph()
        run = p.add_run(f"{heading_text}: ")
        run.bold = True
        process_inline_markdown(p, part)

    doc.add_page_break()

    # ── Body Sections ───────────────────────────────────────────
    body_sections = [
        ("Introduction", "## Introduction"),
        ("Results", "## Results"),
        ("Discussion", "## Discussion"),
        ("Methods", "## Methods"),
        ("Data and Code Availability", "## Data and Code Availability"),
        ("Figure Legends", "## Figure Legends"),
    ]

    processed_headings = {"## Abstract"}  # Already processed

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip frontmatter (title, highlights) and abstract
        if (line.startswith("# ") and not line.startswith("## ")) or line == "## Highlights" or line == "## Abstract":
            # Skip until next major heading
            i += 1
            while i < len(lines) and not (lines[i].strip().startswith("## ") and lines[i].strip() not in ("## Highlights", "## Abstract")):
                i += 1
            continue

        # Section headings
        if line.startswith("## "):
            heading_text = line[3:]
            # Check if it's a major section
            add_heading_styled(doc, heading_text, level=1)
            i += 1
            continue

        if line.startswith("### "):
            heading_text = line[4:]
            add_heading_styled(doc, heading_text, level=2)
            i += 1
            continue

        # Horizontal rule → skip
        if line == "---":
            i += 1
            continue

        # Markdown table
        if line.startswith("|") and i + 2 < len(lines) and lines[i + 1].strip().startswith("|"):
            header, rows, next_i = parse_markdown_table(lines, i)
            add_table_from_markdown(doc, header, rows)
            i = next_i
            continue

        # Bold-only paragraph (subsection marker like **From Literature...**)
        if line.startswith("**") and line.endswith("**") and len(line) > 10:
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
            i += 1
            continue

        # Regular paragraph
        if line:
            add_body_paragraph(doc, line)
        else:
            # Empty line
            doc.add_paragraph("")

        i += 1

    # ── Acknowledgements placeholder ────────────────────────────
    add_heading_styled(doc, "Acknowledgments", level=1)
    doc.add_paragraph("[TO BE COMPLETED]")

    # ── Sources of Funding ──────────────────────────────────────
    add_heading_styled(doc, "Sources of Funding", level=1)
    doc.add_paragraph("[TO BE COMPLETED]")

    # ── Disclosures ─────────────────────────────────────────────
    add_heading_styled(doc, "Disclosures", level=1)
    doc.add_paragraph("None.")

    # ── References ──────────────────────────────────────────────
    ref_path = PROJECT_ROOT / "manuscript" / "references_final.md"
    if ref_path.exists():
        doc.add_page_break()
        add_heading_styled(doc, "References", level=1)

        ref_text = ref_path.read_text(encoding="utf-8")
        # Parse reference entries (numbered lines starting with "N. **...")
        in_refs = False
        for line in ref_text.split("\n"):
            line = line.strip()
            if not line or line.startswith("#") or line.startswith("---") or line.startswith("All references"):
                continue
            if line.startswith("## Final Reference List"):
                continue
            # Reference entry line
            if line[0].isdigit() and ". " in line[:4]:
                p = doc.add_paragraph()
                # Bold the author part up to first period after authors
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Cm(-1.27)
                p.paragraph_format.left_indent = Cm(1.27)
                process_inline_markdown(p, line)
            elif in_refs or (line and line[0].isdigit()):
                # Continuation of previous ref (rare in this format)
                p = doc.add_paragraph()
                process_inline_markdown(p, line)
        print(f"References appended from: {ref_path.name}")
    else:
        print(f"WARNING: Reference file not found: {ref_path}")

    # ── Save ────────────────────────────────────────────────────
    doc.save(str(OUT_PATH))
    print(f"Manuscript saved to: {OUT_PATH}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")


if __name__ == "__main__":
    build_manuscript()
