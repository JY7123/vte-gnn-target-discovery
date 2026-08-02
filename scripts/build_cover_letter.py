#!/usr/bin/env python3
"""Generate ATVB cover letter as Word document."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(r"D:\JY\work\my work\新思路\投稿ATVB")
OUT_DIR.mkdir(parents=True, exist_ok=True)

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# Date
doc.add_paragraph("August 1, 2026")

doc.add_paragraph("")

# Addressee
doc.add_paragraph("The Editors")
doc.add_paragraph("Arteriosclerosis, Thrombosis, and Vascular Biology")
doc.add_paragraph("American Heart Association")

doc.add_paragraph("")

# Salutation
doc.add_paragraph("Dear Editors,")

doc.add_paragraph("")

# Body - Paragraph 1: Submission and manuscript info
p = doc.add_paragraph()
p.add_run(
    "We are pleased to submit our manuscript entitled "
).italic = False
p.add_run(
    '"Data-Leakage-Free Heterogeneous Graph Learning Prioritizes '
    'Cell-Type-Specific Inflammatory and Fibrotic Programs in '
    'Venous Thromboembolism"'
).italic = True
p.add_run(
    " for consideration in Arteriosclerosis, Thrombosis, and Vascular Biology."
)

doc.add_paragraph("")

# Paragraph 2: Background and problem
p = doc.add_paragraph()
p.add_run(
    "Venous thromboembolism (VTE) and its chronic sequela, post-thrombotic "
    "syndrome (PTS), represent a continuous pathological spectrum in which "
    "acute intravascular coagulation progresses to chronic vein wall fibrotic "
    "remodeling. Current therapeutic strategies—anticoagulants targeting "
    "the fluid-phase coagulation cascade—fail to attenuate this structural "
    "remodeling, leaving up to 50% of DVT patients with debilitating PTS. "
    "A fundamental obstacle has been the inability to systematically identify "
    "the molecular drivers of vein wall pathology without being confounded "
    "by the dense literature connectivity of canonical coagulation factors."
)

doc.add_paragraph("")

# Paragraph 3: What we did
p = doc.add_paragraph()
p.add_run(
    "To address this challenge, we developed an integrated multi-scale "
    "framework that combines data-leakage-free heterogeneous graph learning "
    "with primary single-cell transcriptomics and cross-species clinical "
    "validation. We established a rigorous evaluation benchmark (random "
    "stratified 80/10/10 split across 5 independent random seeds, with "
    "per-edge-type deduplication) to evaluate a Tempered Heterogeneous Graph "
    "Transformer (Tempered HGT) against four classical knowledge graph "
    "embedding architectures across 706 million candidate node pairs."
)

doc.add_paragraph("")

# Paragraph 4: Key findings
p = doc.add_paragraph()
p.add_run("Our study produced three main findings:")

doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run(
    "First, under strict data-leakage-free evaluation, Tempered HGT achieved "
    "a Filtered MRR of 0.086 ± 0.029 and Hits@10 of 0.184 ± 0.068, "
    "outperforming the strongest classical baseline (RotatE) by 2.5-fold "
    "in Filtered MRR. Importantly, entity-resolved normalization prioritized "
    "SMAD4—a central transcriptional mediator of TGF-β signaling—as the "
    "top-ranked candidate target (Rank #1), shifting the therapeutic focus "
    "from canonical coagulation factors to solid-phase vein wall remodeling."
)

doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run(
    "Second, single-cell RNA sequencing of 21,230 vein wall cells from a "
    "mouse IVC stenosis model (Day 14) resolved a paracrine axis wherein "
    "infiltrating macrophages serve as the dominant source of Tgfb1 ligand "
    "(mean expression = 1.05), engaging downstream Smad4 transcription "
    "(positive-cell mean expression = 1.19, 41.1% of fibroblasts) and "
    "extracellular matrix execution (Col1a1, Fn1) in adventitial fibroblasts. "
    "This cell-type-specific division of labor provides a concrete mechanistic "
    "basis for post-thrombotic vein wall fibrosis."
)

doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run(
    "Third, cross-species gene set enrichment analysis in a clinical "
    "whole-blood transcriptomic cohort of human VTE patients (GSE48000, "
    "n = 132) confirmed significant enrichment of the vein wall "
    "fibroblast activation program (NES = 1.60, p_empirical = 0.0055 "
    "against 2,000 size-matched random gene sets), demonstrating evolutionary "
    "conservation of this pathological program."
)

doc.add_paragraph("")

# Paragraph 5: Fit for ATVB
p = doc.add_paragraph()
p.add_run(
    "We believe this work is well-suited for the readership of ATVB for "
    "three reasons. First, the discovery of a macrophage-fibroblast TGF-β/"
    "SMAD4 paracrine axis driving post-thrombotic vein wall fibrosis directly "
    "addresses the journal's core focus on vascular biology and thrombosis. "
    "Second, the methodological framework—bridging macro-scale graph "
    "learning with micro-scale single-cell niche mapping—offers a broadly "
    "applicable paradigm for systematic target discovery in cardiovascular "
    "disease, one that actively guards against the data leakage and "
    "literature-density bias that have limited prior computational efforts. "
    "Third, all analyses were conducted with rigorous statistical controls "
    "(leave-one-gene-out sensitivity, empirical permutation testing, 5-seed "
    "reproducibility), and methods are described with full transparency "
    "regarding the evaluation split design."
)

doc.add_paragraph("")

# Paragraph 6: Ethics and declarations
p = doc.add_paragraph()
p.add_run(
    "This manuscript has not been published previously and is not under "
    "consideration elsewhere. All authors have reviewed and approved the "
    "manuscript. The authors declare no competing financial interests. "
    "The complete analysis pipeline and source code are publicly available "
    "at GitHub (https://github.com/JY7123/vte-gnn-target-discovery) with "
    "a permanent Zenodo DOI (10.5281/zenodo.21757258). The mouse IVC "
    "stenosis scRNA-seq data are available from the corresponding author "
    "upon reasonable request."
)

doc.add_paragraph("")

# Paragraph 7: Closing
p = doc.add_paragraph()
p.add_run(
    "We appreciate your consideration of this manuscript and look forward "
    "to your response."
)

doc.add_paragraph("")
doc.add_paragraph("")

# Signature
doc.add_paragraph("Sincerely,")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("Ruihua Wang, M.D., Ph.D.").bold = True
doc.add_paragraph("Department of Vascular Surgery")
doc.add_paragraph("Shanghai Ninth People's Hospital")
doc.add_paragraph("Shanghai Jiao Tong University School of Medicine")
doc.add_paragraph("Shanghai, China")
doc.add_paragraph("Email: wangruihua@sjtu.edu.cn")
doc.add_paragraph("Phone: 19121035097")

# Save
OUT_PATH = OUT_DIR / "cover_letter.docx"
doc.save(str(OUT_PATH))
print(f"Cover letter saved to: {OUT_PATH}")
