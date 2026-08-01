#!/usr/bin/env python3
"""
Build ATVB-formatted Supplementary Materials docx from markdown source.

Converts manuscript/supplementary_tables.md → 投稿ATVB/Supplementary_Materials.docx
reusing the formatting conventions of build_atvb_manuscript.py
(Times New Roman 12 pt, double spacing, Table Grid with gray header rows).
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Reuse the shared markdown→docx helpers from the manuscript builder.
from build_atvb_manuscript import (
    add_body_paragraph,
    add_heading_styled,
    add_table_from_markdown,
    parse_markdown_table,
)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "manuscript" / "supplementary_tables.md"
MANUSCRIPT_TITLE = (
    "Temporally Evaluated Heterogeneous Graph Learning Prioritizes "
    "Cell-Type-Specific Inflammatory and Fibrotic Programs in Venous Thromboembolism"
)
OUT_DIR = PROJECT_ROOT.parent / "投稿ATVB"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "Supplementary_Materials.docx"


def add_title_block(doc):
    """Centered title + manuscript subtitle, then a page break."""
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Supplementary Materials")
    title_run.bold = True
    title_run.font.size = Pt(14)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(MANUSCRIPT_TITLE)
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    doc.add_paragraph("")
    doc.add_page_break()


def build_supplementary():
    if not MD_PATH.exists():
        print(f"ERROR: {MD_PATH} not found")
        sys.exit(1)

    lines = MD_PATH.read_text(encoding="utf-8").split("\n")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0

    add_title_block(doc)

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip the markdown title + subtitle at the top
        if line.startswith("# ") or line.startswith("**Manuscript:"):
            i += 1
            continue

        # Horizontal rule → skip
        if line == "---":
            i += 1
            continue

        # Section heading (## Table S#: ...)
        if line.startswith("## "):
            add_heading_styled(doc, line[3:], level=2)
            i += 1
            continue

        # Markdown table
        if line.startswith("|") and i + 2 < len(lines) and lines[i + 1].strip().startswith("|"):
            header, rows, next_i = parse_markdown_table(lines, i)
            add_table_from_markdown(doc, header, rows)
            i = next_i
            continue

        # Source/notes paragraph starting with "Source:"
        if line.startswith("Source:"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.italic = True
            i += 1
            continue

        # Regular paragraph
        if line:
            add_body_paragraph(doc, line)
        else:
            doc.add_paragraph("")

        i += 1

    doc.save(str(OUT_PATH))
    print(f"Supplementary Materials saved to: {OUT_PATH}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")


if __name__ == "__main__":
    build_supplementary()
